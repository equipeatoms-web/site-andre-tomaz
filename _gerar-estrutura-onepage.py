# -*- coding: utf-8 -*-
"""Gera a estrutura da one page do Andre Tomaz em DOCX (documento de discussao).
Rodar: python _gerar-estrutura-onepage.py  ->  estrutura-onepage.docx"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OBSIDIAN = RGBColor(0x14, 0x14, 0x1A)
AMBER = RGBColor(0xE8, 0x9B, 0x3D)
GRAY = RGBColor(0x6B, 0x72, 0x7A)
HEAD_FONT = "Space Grotesk"
BODY_FONT = "Inter"

doc = Document()
for s in doc.sections:
    s.top_margin = Cm(2.0); s.bottom_margin = Cm(2.0)
    s.left_margin = Cm(2.2); s.right_margin = Cm(2.2)

n = doc.styles["Normal"]
n.font.name = BODY_FONT; n.font.size = Pt(10.5); n.font.color.rgb = OBSIDIAN
n.element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
n.paragraph_format.space_after = Pt(6); n.paragraph_format.line_spacing = 1.15


def sf(run, name=BODY_FONT, size=10.5, bold=False, italic=False, color=OBSIDIAN):
    run.font.name = name
    run.element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size); run.font.bold = bold
    run.font.italic = italic; run.font.color.rgb = color


def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20); p.paragraph_format.space_after = Pt(4)
    sf(p.add_run(text), name=HEAD_FONT, size=15, bold=True)
    pPr = p._p.get_or_add_pPr(); pBdr = OxmlElement("w:pBdr"); b = OxmlElement("w:bottom")
    b.set(qn("w:val"), "single"); b.set(qn("w:sz"), "8"); b.set(qn("w:space"), "2"); b.set(qn("w:color"), "E89B3D")
    pBdr.append(b); pPr.append(pBdr)


def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(3)
    sf(p.add_run(text), name=HEAD_FONT, size=11.5, bold=True, color=AMBER)


def para(text, size=10.5, italic=False, color=OBSIDIAN, after=6, bold=False):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(after)
    sf(p.add_run(text), size=size, italic=italic, color=color, bold=bold)
    return p


def rot(label, text):
    """Linha com rotulo em ambar."""
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(3)
    sf(p.add_run(label + " "), size=9.5, bold=True, color=AMBER)
    sf(p.add_run(text), size=10.5)


def bullet(lead, text=""):
    p = doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after = Pt(3)
    if lead: sf(p.add_run(lead + (" " if text else "")), bold=True)
    if text: sf(p.add_run(text))


def box(text, label=None):
    """Texto proposto da pagina."""
    if label:
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
        sf(p.add_run(label), size=9, bold=True, color=GRAY)
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    cell = t.cell(0, 0)
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), "F4F1EA")
    cell._tc.get_or_add_tcPr().append(shd)
    first = True
    for ln in text.split("\n"):
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        p.paragraph_format.space_after = Pt(4)
        if ln.startswith("##"):
            sf(p.add_run(ln[2:].strip()), name=HEAD_FONT, size=14, bold=True)
        elif ln.startswith("#"):
            sf(p.add_run(ln[1:].strip()), name=HEAD_FONT, size=11.5, bold=True, color=AMBER)
        else:
            sf(p.add_run(ln), size=10.5)
    doc.add_paragraph()


def tabela(rows, header, widths):
    t = doc.add_table(rows=0, cols=len(header)); t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    r = t.add_row()
    for i, hd in enumerate(header):
        sf(r.cells[i].paragraphs[0].add_run(hd), bold=True, size=9.5, color=AMBER)
    for row in rows:
        r = t.add_row()
        for i, cellv in enumerate(row):
            sf(r.cells[i].paragraphs[0].add_run(cellv), size=9.5, bold=(i == 0))
            r.cells[i].width = Cm(widths[i])
    doc.add_paragraph()


# ============ CAPA ============
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(4)
sf(p.add_run("ONE PAGE"), name=HEAD_FONT, size=26, bold=True)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sf(p.add_run("André Tomaz · Consultor de Transformação Empresarial"), name=HEAD_FONT, size=12, color=AMBER)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sf(p.add_run("Estrutura para discussão · 10 de agosto de 2026 · textos propostos, não finais"), size=9, italic=True, color=GRAY)
doc.add_paragraph()

para("Uma one page não é um site institucional reduzido. É uma máquina de qualificação: em 3 minutos de leitura ela precisa fazer o visitante se reconhecer no problema, entender que existe solução, confiar que você já fez isso antes e dar o próximo passo sozinho.", italic=True)
para("A régua desta página: cada seção responde a uma pergunta que o leitor faz nessa ordem exata. Se alguma seção não responder a uma pergunta real, ela sai.", italic=True)

# ============ ARQUITETURA ============
h1("A arquitetura em 9 blocos")
tabela([
    ("1. Abertura", "\u201CIsso é para mim?\u201D", "Reconhecimento em 5 segundos"),
    ("2. O espelho", "\u201CEle entende meu problema?\u201D", "Nomear a dor melhor que o próprio leitor"),
    ("3. Como eu trabalho", "\u201CO que ele faz exatamente?\u201D", "Clareza sem jargão"),
    ("4. O método", "\u201CComo funciona na prática?\u201D", "Previsibilidade, tira o medo do desconhecido"),
    ("5. O que se constrói", "\u201CO que eu recebo?\u201D", "Concretude sem virar cardápio"),
    ("6. O caso real", "\u201CJá funcionou com alguém?\u201D", "Prova"),
    ("7. Para quem é", "\u201CSirvo para isso?\u201D", "Qualificar e filtrar (autoridade)"),
    ("8. A porta", "\u201CE agora?\u201D", "Um único próximo passo, de baixo atrito"),
    ("9. Perguntas", "\u201CMas e se…?\u201D", "Derrubar a objeção antes do silêncio"),
], ("Bloco", "A pergunta do leitor", "O que precisa entregar"), (3.6, 5.4, 7.2))

para("Ordem inegociável: o problema vem antes da solução, e a prova vem antes da porta. Página que abre falando de si mesma perde o leitor no primeiro scroll.", bold=True)

doc.add_page_break()

# ============ BLOCO 1 ============
h1("Bloco 1 · Abertura (hero)")
rot("Objetivo:", "o visitante se reconhece e entende o que você faz, sem rolar a página.")
rot("Estrutura:", "uma frase-problema grande, uma linha explicando quem você é, e um caminho discreto para a porta.")
box("## Sua empresa não precisa de mais funcionários. Precisa de um ecossistema digital inteligente.\n\nConsultor de transformação empresarial com foco em processos, sistemas e inteligência artificial.\n\n[ Começar pelo diagnóstico ]", label="TEXTO PROPOSTO:")
h2("Por que essa frase")
para("Ela já é a sua bio no Instagram, já foi testada, e faz o trabalho mais difícil de um hero: fala do problema do leitor (a tentação de contratar mais gente) e não do seu currículo. O visitante se vê na primeira linha.")
h2("As alternativas, se você quiser testar")
bullet("A.", "\u201CPrimeiro eu entendo como a empresa funciona. Depois eu construo o que ela precisa.\u201D (a frase-mãe da marca; mais calma, menos provocativa)")
bullet("B.", "\u201CA sua empresa continua funcionando quando você não está?\u201D (pergunta-espelho; forte, mas arriscada como primeira linha porque exige que o leitor admita algo antes de confiar em você)")
rot("Decisão pendente:", "vídeo de 30 segundos no hero ou só foto? Um vídeo seu falando os 30 segundos da narrativa aumenta muito a conversão em serviço de alto valor.")

# ============ BLOCO 2 ============
h1("Bloco 2 · O espelho (o problema)")
rot("Objetivo:", "descrever a rotina dele com tanta precisão que ele pensa \u201Cesse cara já esteve na minha empresa\u201D.")
rot("Formato:", "lista curta de sintomas, primeira pessoa do leitor. Sem adjetivo, só fato.")
box("# Talvez você reconheça isso\n\nA operação anda enquanto você está por perto. Quando você viaja, atrasa.\n\nA equipe erra em coisas que você achava resolvidas: preço trocado, horário sobreposto, cliente sem retorno.\n\nO cliente manda mensagem às 23h e você responde no dia seguinte, quando ele já procurou outro lugar.\n\nA informação existe, mas está espalhada entre planilha, caderno, WhatsApp e a memória de duas pessoas.\n\nVocê já pensou em contratar mais alguém. E também já percebeu que mais gente, sem estrutura, é mais gente para gerenciar.", label="TEXTO PROPOSTO:")
rot("Regra:", "todo sintoma aqui tem que ter saído de uma empresa real que você viu. Sintoma inventado soa a copy de agência e o leitor sente.")

# ============ BLOCO 3 ============
h1("Bloco 3 · Como eu trabalho")
rot("Objetivo:", "explicar o serviço em linguagem de dono, sem jargão. Esse é o bloco que você já escreveu.")
box("# O que eu faço\n\nEu entro na empresa para entender como ela funciona hoje. Analiso operação, atendimento, processos e gestão. A partir disso desenho soluções que organizam o negócio e eliminam gargalos.\n\nEm alguns casos isso vira um sistema próprio. Em outros, uma IA, automações ou uma nova estrutura de operação.\n\nCada empresa precisa de uma solução diferente.", label="TEXTO PROPOSTO (o parágrafo oficial, sem alteração):")
h2("O contraste (opcional, mas poderoso aqui)")
box("Consultor tradicional entrega diagnóstico.\nSoftware house entrega sistema.\nEu entrego a operação rodando.", label="BLOCO DE CONTRASTE:")
para("Esse trio posiciona você contra as duas alternativas que o leitor já considerou. Recomendo manter: é a frase que mais rápido explica o seu diferencial.")

# ============ BLOCO 4 ============
h1("Bloco 4 · O método")
rot("Objetivo:", "mostrar que existe um caminho definido. Previsibilidade vende serviço caro.")
rot("Formato:", "linha do tempo horizontal (desktop) ou vertical (celular), com as 8 etapas em uma linha cada.")
box("# Como funciona\n\n0. Protocolo de marca. Um diagnóstico on-line da sua operação e da sua presença digital. É por aqui que todo trabalho começa.\n1. Imersão. Eu entro na empresa e observo como ela funciona de verdade.\n2. Diagnóstico. Onde a empresa perde tempo, dinheiro e cliente.\n3. Arquitetura. O desenho de como a operação deveria funcionar.\n4. Priorização. O que gera mais impacto vem primeiro.\n5. Desenvolvimento. A construção sob medida: sistema, IA, automações ou nova estrutura.\n6. Implantação. Colocar para rodar na operação real, com a equipe usando.\n7. Acompanhamento. Medir, ajustar e evoluir junto com o negócio.", label="TEXTO PROPOSTO:")
rot("Detalhe que vende:", "destacar visualmente que as etapas 1 a 3 acontecem ANTES de qualquer linha de código. É o que separa você da software house.")

# ============ BLOCO 5 ============
doc.add_page_break()
h1("Bloco 5 · O que pode ser construído")
rot("Objetivo:", "dar concretude. O leitor precisa visualizar o que existe no fim.")
rot("Cuidado:", "isso não é um cardápio de produtos. Abre e fecha com \u201Ccada empresa precisa de uma solução diferente\u201D para não virar tabela de preços mental.")
box("# O que costuma nascer disso\n\nPainel de gestão. A operação inteira em uma tela: agenda, cadastro, preços por unidade, pacotes, histórico de quem fez o quê.\n\nAtendimento com IA. Um agente que atende, agenda, qualifica e vende no WhatsApp, em vários idiomas, sem horário para começar ou terminar.\n\nPlataforma de conhecimento. Quando o que a empresa sabe vira produto: curso, formação, comunidade.\n\nPresença digital. Site e conteúdo que atraem o cliente certo e sustentam o preço.\n\nInteligência de dados. Números em tempo real e alerta quando algo sai do planejado.\n\nNenhuma empresa recebe tudo isso. Cada uma recebe o que resolve o gargalo dela.", label="TEXTO PROPOSTO:")

# ============ BLOCO 6 ============
h1("Bloco 6 · O caso real")
rot("Objetivo:", "provar. Sem prova, tudo acima é promessa.")
rot("⚠️ Trava:", "só afirmar o que foi ENTREGUE. Nenhum número de resultado até as métricas do painel serem extraídas. Nomear a cliente exige autorização por escrito.")
box("# Um instituto premium na Bélgica\n\nA fundadora era o centro de tudo: o atendimento, a agenda, o padrão da marca, a decisão de cada caso. O negócio crescia e ela virou o limite dele.\n\nEntrei para entender a operação inteira antes de construir qualquer coisa. O que nasceu disso:\n\nUm painel de gestão que guia a recepção, cobra o preço certo por unidade e controla pacotes de sessão.\n\nUma atendente digital no WhatsApp que agenda em quatro idiomas, 24 horas por dia, preenche a ficha de saúde antes da visita e registra por onde cada cliente chegou.\n\nUma plataforma própria de formação, com uma agente comercial que conduz a matrícula do início ao fim.\n\nUm site em quatro idiomas, construído dentro das regras do setor de saúde na Bélgica.\n\nHoje a dona ajusta preço, tom de voz e regras do atendimento sozinha, sem programador.", label="TEXTO PROPOSTO (versão sem nomear a cliente):")
h2("O que falta para esse bloco ficar imbatível")
bullet("Métricas do painel:", "quantos agendamentos a agente fez, quantos fora do horário comercial, percentual agente contra recepção. O painel já registra tudo isso.")
bullet("Depoimento em vídeo:", "as 8 perguntas de entrevista já estão prontas na narrativa. Com termo assinado, o rosto dela nesta página vale mais que qualquer texto.")
bullet("Autorização de nome:", "enquanto não houver, \u201Cum instituto premium na Bélgica\u201D funciona e ainda cria curiosidade.")

# ============ BLOCO 7 ============
h1("Bloco 7 · Para quem é (e para quem não é)")
rot("Objetivo:", "qualificar. Dizer para quem você não serve é o que mais aumenta a confiança de quem serve.")
box("# Esse trabalho é para você se\n\nVocê tem um negócio que já funciona e já tem demanda. O problema não é falta de cliente, é a operação não acompanhar.\n\nA sua marca depende da sua presença e do seu nome.\n\nVocê atua em um mercado onde uma palavra errada tem consequência legal.\n\nVocê quer transformar o que sabe em produto, não só em atendimento.\n\n# Não é para você se\n\nVocê procura tráfego pago, posts ou gestão de redes sociais. Isso é outro serviço, com outro profissional.\n\nVocê quer o mais barato. Sob medida não compete com pacote pronto.\n\nVocê quer entregar o problema e não participar. As três primeiras etapas exigem você dentro.", label="TEXTO PROPOSTO:")

# ============ BLOCO 8 ============
h1("Bloco 8 · A porta")
rot("Objetivo:", "um único próximo passo. Duas portas é o mesmo que nenhuma.")
rot("A porta escolhida:", "o Protocolo de Marca (etapa 0). Baixo compromisso para ele, qualificação para você.")
box("# Comece pelo diagnóstico\n\nAntes de qualquer proposta, eu faço uma leitura da sua operação e da sua presença digital. Você recebe um retrato honesto de onde a empresa está e onde ela perde.\n\nSe fizer sentido seguir, a gente conversa. Se não fizer, o retrato é seu do mesmo jeito.\n\n[ Solicitar o diagnóstico ]", label="TEXTO PROPOSTO:")
h2("Coerente com a doutrina de CTA")
para("A página inteira não pede clique em nenhum outro lugar. Aqui, no fim, existe uma porta única e sóbria. Não é \u201Cme chama\u201D: é uma oferta de valor que a pessoa aceita se quiser. Ela chegou até aqui rolando a página inteira, ou seja, já está interessada.")
rot("Decisões pendentes:", "o diagnóstico é gratuito ou pago? Formulário de 4 campos ou WhatsApp direto? Meu voto: formulário curto (nome, empresa, setor, o maior gargalo hoje), porque já qualifica e evita conversa fria.")

# ============ BLOCO 9 ============
h1("Bloco 9 · Perguntas que travam a decisão")
rot("Objetivo:", "derrubar a objeção no texto, porque ela não vai ser dita em voz alta.")
box("Quanto tempo leva? Depende do tamanho do gargalo. O diagnóstico é rápido. A construção é definida na etapa de priorização, com prazo fechado antes de começar.\n\nEu vou ficar dependente de você? Não. Tudo que eu construo é feito para você operar sozinho. No caso da Bélgica, a dona ajusta preço, tom e regras sem me chamar.\n\nJá tenho um sistema. Ótimo. Muitas vezes o problema não é o sistema, é o desenho em volta dele. O diagnóstico mostra se vale integrar, corrigir ou substituir.\n\nMinha empresa é pequena demais? Se você é o gargalo, o tamanho não importa. O que importa é se a operação já tem repetição suficiente para ser estruturada.\n\nE os meus dados? Estrutura de acesso, separação de ambiente e as regras do seu setor entram no desenho desde a primeira etapa, não como remendo no fim.", label="TEXTO PROPOSTO:")

# ============ VISAO ACE ============
doc.add_page_break()
h1("👁️ VISÃO ACE aplicada à página")
bullet("Nível de consciência.", "O ICP é consciente do problema (\u201Ceu sou o gargalo\u201D) e inconsciente da solução: ele só conhece agência, chatbot e contratar mais gente. Por isso a página precisa NOMEAR a categoria antes de vender. É o que os blocos 2 e 3 fazem: primeiro o espelho, depois \u201Cexiste um jeito de estruturar isso\u201D.")
bullet("Escada de valor.", "A página vende só o primeiro degrau (o diagnóstico). Tentar vender a implantação inteira aqui mata a conversão: ticket alto não se vende em página, se vende em conversa. A página existe para gerar a conversa certa.")
bullet("Densidade do funil.", "O bloco 7 (para quem não é) parece que afasta, mas é o que mais qualifica. Menos leads, muito melhores. Para um consultor que atende poucos clientes por vez, isso é ganho puro.")
bullet("Prova antes da porta.", "O bloco 6 precisa estar acima do bloco 8. Ninguém aceita um diagnóstico de quem não provou nada.")
bullet("O risco a evitar.", "Página com números inventados converte no curto prazo e destrói autoridade no médio. Enquanto as métricas do case não forem extraídas, a página fala de entregas. Assim que existirem, elas entram no bloco 6 e a conversão sobe de novo, agora com lastro.")

# ============ DECISOES ============
h1("Decisões que preciso de você para construir")
tabela([
    ("Domínio", "andretomaz.com? Outro? Já existe algo publicado nele hoje?"),
    ("Hospedagem", "VPS com EasyPanel (você já tem) ou algo mais simples? Recomendo o seu VPS: custo zero adicional e você controla."),
    ("Idioma", "Só português ou português e inglês? Você está circulando na Europa e o case é belga. Duas versões custam pouco se a estrutura já nascer preparada."),
    ("O diagnóstico", "Gratuito ou pago? Nome definitivo (\u201CProtocolo de Marca\u201D fica ou muda?)"),
    ("Contato", "Formulário curto, WhatsApp, ou agenda com horário? Meu voto: formulário de 4 campos."),
    ("Vídeo no topo", "Você grava 30 segundos falando a narrativa? Aumenta muito a conversão."),
    ("O nome da cliente", "Podemos citar Maison Ipanema e Neiva Cimini, ou fica \u201Cum instituto premium na Bélgica\u201D por enquanto?"),
    ("Foto nova", "As fotos de estúdio que você já tem servem, ou vale uma sessão nova alinhada à identidade?"),
], ("Decisão", "A pergunta"), (4.2, 12.0))

h1("Como eu construo depois de aprovada")
bullet("1.", "Uma página HTML única, com a sua identidade (obsidiana e âmbar, Space Grotesk e Inter), leve e rápida.")
bullet("2.", "Todo texto e cor em um arquivo de configuração separado, para você trocar sem mexer no código, e para a estrutura servir de base a outros clientes depois.")
bullet("3.", "Preparada para ser lida por IA (dados estruturados), como fizemos no site da Bélgica.")
bullet("4.", "Responsiva de verdade: mais da metade vai ler no celular, vindo do Instagram.")
bullet("5.", "Com medição desde o primeiro dia, para saber quantos leem, até onde rolam e quantos pedem o diagnóstico.")

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(14)
sf(p.add_run("Fonte deste documento: G:\\ANDRE TOMAZ\\site\\01-estrutura-onepage.md"), size=8.5, italic=True, color=GRAY)

OUT = "estrutura-onepage.docx"
doc.save(OUT)
print("OK:", OUT)
